"""
GuíaBot — Extractor de texto .docx
Extrae contenido, cuenta palabras y párrafos para análisis CAC
"""

import re
from docx import Document
import io


def extraer_texto_docx(contenido_bytes: bytes) -> dict:
    """
    Extrae texto completo de un .docx y retorna métricas básicas.

    Retorna:
        texto_completo: str
        parrafos: list[str]
        conteo_palabras: int
        conteo_parrafos: int
        tiene_tablas: bool
        texto_tablas: str
    """
    doc = Document(io.BytesIO(contenido_bytes))

    parrafos = []
    for p in doc.paragraphs:
        texto = p.text.strip()
        if texto:
            parrafos.append(texto)

    # Extraer texto de tablas también
    texto_tablas = ""
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    texto_tablas += celda.text.strip() + " "

    texto_completo = "\n\n".join(parrafos)
    if texto_tablas:
        texto_completo += "\n\n[CONTENIDO DE TABLAS]\n" + texto_tablas

    # Contar palabras (ignorar líneas muy cortas como encabezados APA)
    palabras = len(re.findall(r'\b\w+\b', texto_completo))

    return {
        "texto_completo": texto_completo,
        "parrafos": parrafos,
        "conteo_palabras": palabras,
        "conteo_parrafos": len(parrafos),
        "tiene_tablas": len(doc.tables) > 0,
        "texto_tablas": texto_tablas,
    }


def verificar_escala_cac(grado: str, conteo_palabras: int, conteo_parrafos: int) -> dict:
    """
    Verifica si el texto cumple la escala progresiva CAC según el grado.
    """
    ESCALA = {
        "6":  {"palabras_min": 80,  "palabras_max": 150,  "parrafos_min": 1, "parrafos_max": 2},
        "7":  {"palabras_min": 150, "palabras_max": 250,  "parrafos_min": 2, "parrafos_max": 3},
        "8":  {"palabras_min": 250, "palabras_max": 350,  "parrafos_min": 3, "parrafos_max": 4},
        "9":  {"palabras_min": 350, "palabras_max": 500,  "parrafos_min": 4, "parrafos_max": 5},
        "10": {"palabras_min": 500, "palabras_max": 700,  "parrafos_min": 5, "parrafos_max": 7},
        "11": {"palabras_min": 800, "palabras_max": 1200, "parrafos_min": 8, "parrafos_max": 12},
    }

    escala = ESCALA.get(str(grado))
    if not escala:
        return {"cumple": None, "razon": f"Grado '{grado}' no está en la escala CAC."}

    cumple_palabras = escala["palabras_min"] <= conteo_palabras <= escala["palabras_max"]
    cumple_parrafos = escala["parrafos_min"] <= conteo_parrafos <= escala["parrafos_max"]

    observaciones = []
    if not cumple_palabras:
        if conteo_palabras < escala["palabras_min"]:
            observaciones.append(
                f"Palabras insuficientes: {conteo_palabras} (mínimo {escala['palabras_min']})"
            )
        else:
            observaciones.append(
                f"Palabras excesivas: {conteo_palabras} (máximo {escala['palabras_max']})"
            )

    if not cumple_parrafos:
        if conteo_parrafos < escala["parrafos_min"]:
            observaciones.append(
                f"Párrafos insuficientes: {conteo_parrafos} (mínimo {escala['parrafos_min']})"
            )
        else:
            observaciones.append(
                f"Párrafos en exceso: {conteo_parrafos} (máximo {escala['parrafos_max']})"
            )

    return {
        "cumple": cumple_palabras and cumple_parrafos,
        "conteo_palabras": conteo_palabras,
        "conteo_parrafos": conteo_parrafos,
        "esperado_palabras": f"{escala['palabras_min']}–{escala['palabras_max']}",
        "esperado_parrafos": f"{escala['parrafos_min']}–{escala['parrafos_max']}",
        "observaciones": observaciones,
    }
